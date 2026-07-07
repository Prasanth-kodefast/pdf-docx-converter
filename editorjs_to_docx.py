"""
editorjs_to_docx
================

Render a CLM contract (per-page EditorJS blocks + per-page headers /
footers) into a real `.docx` file where the sender's redline marks
become native Word Track Changes:

  <ins class="clm-suggestion-insert">TEXT</ins>  →  <w:ins ...><w:r><w:t>TEXT
  <mark class="clm-redline">TEXT</mark>          →  <w:del ...><w:r><w:delText>TEXT

Once the counterparty opens the file in Microsoft Word (or LibreOffice
or Google Docs), the Review ribbon lights up and they can accept /
reject / add further changes exactly the same way they would with any
tracked-changes document from a peer.

This is the enterprise-standard flow. Ironclad + DocuSign CLM Negotiate
do the same round-trip; the counterparty never needs to log into a web
app or install a Word plugin — they just work in Word.

Public entry point:

    render_contract_docx(pages, meta) -> bytes

`pages` is a list shaped `[{ pageNo, blocks }]` where `blocks` are raw
EditorJS block objects.  `meta` carries the title, sender identity for
`w:author`, an optional cover message, and optional per-page header /
footer maps keyed by page index.

Returns the assembled DOCX as an in-memory bytestring so the FastAPI
handler can stream it back to the caller.
"""

from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from typing import Any, Dict, Iterable, List, Optional

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


# ─── OOXML namespaces ───────────────────────────────────────────────────────
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ─── Revision id counter ────────────────────────────────────────────────────
class _RevisionCounter:
    """Word requires every `w:ins` / `w:del` to carry a unique `w:id`.
    A shared counter across a whole document keeps ids monotonic, which
    is what every serializer we've inspected (Word, LibreOffice, Google
    Docs) produces."""

    def __init__(self, start: int = 1) -> None:
        self._n = start - 1

    def next(self) -> str:
        self._n += 1
        return str(self._n)


# ─── Inline-run parser ──────────────────────────────────────────────────────
#
# Given a paragraph's inline HTML (as produced by EditorJS with the
# sender's redline marks + mentions + clause chips baked in), walk it
# and emit a stream of typed segments the paragraph renderer knows how
# to attach to a `w:p` element:
#
#   { "kind": "text", "text": str, "bold": bool, "italic": bool,
#     "underline": bool, "strike": bool, "color": Optional[str] }
#   { "kind": "ins",   "runs": [text-run-dicts] }
#   { "kind": "del",   "runs": [text-run-dicts] }
#   { "kind": "linebreak" }
#
# The parser is deliberately depth-first + accumulates the current
# formatting state as it descends, so nested `<strong><ins>foo</ins></strong>`
# still lands in Word as bold + tracked insertion.


_TRACKED_INS_CLASS = "clm-suggestion-insert"
_TRACKED_DEL_CLASS = "clm-redline"

# Elements we treat as "not visible content" — mentions and chip carriers.
# Their visible text is what the counterparty should see; the HTML plumbing
# around them is stripped.
_CHIP_CLASSES = {
    "mention",
    "clause-var-chip",
    "clause-tag-chip",  # hidden carrier — never emit any text for these
}


def _class_list(tag) -> List[str]:
    cls = tag.get("class") if hasattr(tag, "get") else None
    if not cls:
        return []
    if isinstance(cls, str):
        return cls.split()
    return list(cls)


def _has_class(tag, name: str) -> bool:
    return name in _class_list(tag)


def _style_dict(tag) -> Dict[str, str]:
    raw = tag.get("style") if hasattr(tag, "get") else None
    if not raw:
        return {}
    out: Dict[str, str] = {}
    for chunk in raw.split(";"):
        if ":" not in chunk:
            continue
        k, v = chunk.split(":", 1)
        out[k.strip().lower()] = v.strip()
    return out


def _hex_from_color(css: str) -> Optional[str]:
    """Coerce a CSS color value (`#rrggbb`, `rgb(r,g,b)`, name) into a
    6-char hex string usable by python-docx's `RGBColor`. Returns
    None when the color can't be parsed — the run just falls back to
    the paragraph default."""
    if not css:
        return None
    css = css.strip()
    if css.startswith("#") and (len(css) == 7 or len(css) == 4):
        return css[1:].upper() if len(css) == 7 else "".join(c * 2 for c in css[1:]).upper()
    m = re.match(r"rgb\((\d+)\s*,\s*(\d+)\s*,\s*(\d+)\)", css)
    if m:
        r, g, b = (int(x) for x in m.groups())
        return f"{r:02X}{g:02X}{b:02X}"
    return None


def _parse_inline(
    html: str,
) -> List[Dict[str, Any]]:
    """Walk the inline HTML of a block and return a flat sequence of
    typed segments (text runs, ins groups, del groups, line breaks).

    The returned list is what `_apply_runs_to_paragraph` consumes; the
    two functions are the sole boundary between "HTML I don't want
    inside my Word renderer" and "runs I know how to attach."
    """
    if not html:
        return []
    soup = BeautifulSoup(html, "html.parser")
    segments: List[Dict[str, Any]] = []

    def walk(node, fmt: Dict[str, Any]):
        # Text node — emit a run.
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                segments.append({"kind": "text", "text": text, **fmt})
            return

        # Line break.
        if getattr(node, "name", None) == "br":
            segments.append({"kind": "linebreak"})
            return

        # Tracked insertion / deletion — recurse with a marker so the
        # parent emits `w:ins` / `w:del` around the collected runs.
        cls = _class_list(node)
        if node.name == "ins" and _TRACKED_INS_CLASS in cls:
            child_runs: List[Dict[str, Any]] = []
            _walk_into(node, fmt, into=child_runs)
            segments.append({"kind": "ins", "runs": child_runs})
            return
        if node.name in ("mark", "del") and _TRACKED_DEL_CLASS in cls:
            child_runs: List[Dict[str, Any]] = []
            _walk_into(node, fmt, into=child_runs)
            segments.append({"kind": "del", "runs": child_runs})
            return

        # Chip carriers — hidden cite tag we always drop.
        if node.name == "cite" and _has_class(node, "clause-tag-chip"):
            return

        # Mention / clause-var-chip — pass through the visible text,
        # ignore the wrapping element. The value the counterparty sees
        # is whatever text-content the chip renders.
        if _has_class(node, "mention") or _has_class(node, "clause-var-chip"):
            for child in node.children:
                walk(child, fmt)
            return

        # Inline formatting — accumulate onto `fmt` and recurse.
        next_fmt = dict(fmt)
        if node.name in ("strong", "b"):
            next_fmt["bold"] = True
        if node.name in ("em", "i"):
            next_fmt["italic"] = True
        if node.name == "u":
            next_fmt["underline"] = True
        if node.name == "s":
            next_fmt["strike"] = True
        if node.name == "sup":
            next_fmt["superscript"] = True
        if node.name == "sub":
            next_fmt["subscript"] = True
        if node.name == "a":
            next_fmt["hyperlink"] = node.get("href")

        style = _style_dict(node)
        if "color" in style:
            hex_c = _hex_from_color(style["color"])
            if hex_c:
                next_fmt["color"] = hex_c
        if style.get("font-weight") in ("bold", "700", "800", "900"):
            next_fmt["bold"] = True
        if style.get("font-style") == "italic":
            next_fmt["italic"] = True
        if style.get("text-decoration") == "underline":
            next_fmt["underline"] = True
        if style.get("text-decoration") == "line-through":
            next_fmt["strike"] = True

        for child in node.children:
            walk(child, next_fmt)

    def _walk_into(container, fmt: Dict[str, Any], into: List[Dict[str, Any]]):
        nonlocal segments
        saved = segments
        segments = into
        try:
            for child in container.children:
                walk(child, fmt)
        finally:
            segments = saved

    for child in soup.children:
        walk(child, {})
    return segments


# ─── OOXML helpers ──────────────────────────────────────────────────────────
def _make_rPr(run_props: Dict[str, Any]) -> Optional[OxmlElement]:
    """Build a `<w:rPr>` element carrying the formatting for one text
    run. Returns None when the run has no formatting overrides — the
    caller can then just omit rPr entirely, matching what Word does."""
    if not any(run_props.get(k) for k in ("bold", "italic", "underline", "strike", "color", "superscript", "subscript")):
        return None
    rpr = OxmlElement("w:rPr")
    if run_props.get("bold"):
        b = OxmlElement("w:b")
        rpr.append(b)
    if run_props.get("italic"):
        i = OxmlElement("w:i")
        rpr.append(i)
    if run_props.get("underline"):
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)
    if run_props.get("strike"):
        s = OxmlElement("w:strike")
        rpr.append(s)
    if run_props.get("superscript"):
        v = OxmlElement("w:vertAlign")
        v.set(qn("w:val"), "superscript")
        rpr.append(v)
    if run_props.get("subscript"):
        v = OxmlElement("w:vertAlign")
        v.set(qn("w:val"), "subscript")
        rpr.append(v)
    if run_props.get("color"):
        c = OxmlElement("w:color")
        c.set(qn("w:val"), run_props["color"])
        rpr.append(c)
    return rpr


def _make_run(seg: Dict[str, Any], *, deleted: bool = False) -> OxmlElement:
    """Build a `<w:r>` (run) element. When `deleted=True` the text
    child is `<w:delText>` instead of `<w:t>` — Word requires this
    inside `<w:del>` revisions."""
    r = OxmlElement("w:r")
    rpr = _make_rPr(seg)
    if rpr is not None:
        r.append(rpr)
    text_tag = "w:delText" if deleted else "w:t"
    t = OxmlElement(text_tag)
    t.text = seg.get("text", "")
    # Preserve leading / trailing whitespace exactly as authored —
    # Word collapses it otherwise. This mirrors how Word itself
    # serialises runs it emits from the Review ribbon.
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _make_linebreak_run() -> OxmlElement:
    r = OxmlElement("w:r")
    r.append(OxmlElement("w:br"))
    return r


def _make_ins(runs: List[Dict[str, Any]], author: str, when: str, rev: _RevisionCounter) -> OxmlElement:
    """Build a `<w:ins>` (tracked insertion) element wrapping the given
    runs. Word's Review ribbon reads `w:author` + `w:date` to attribute
    the change and the timestamp shown on hover."""
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), rev.next())
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), when)
    for seg in runs:
        if seg.get("kind") == "linebreak":
            ins.append(_make_linebreak_run())
        else:
            ins.append(_make_run(seg))
    return ins


def _make_del(runs: List[Dict[str, Any]], author: str, when: str, rev: _RevisionCounter) -> OxmlElement:
    """Build a `<w:del>` (tracked deletion) element. Deletions use
    `<w:delText>` for the text child; if you accidentally emit `<w:t>`
    inside `w:del` Word will still open the file but won't recognise
    the deletion as a real revision."""
    d = OxmlElement("w:del")
    d.set(qn("w:id"), rev.next())
    d.set(qn("w:author"), author)
    d.set(qn("w:date"), when)
    for seg in runs:
        if seg.get("kind") == "linebreak":
            d.append(_make_linebreak_run())
        else:
            d.append(_make_run(seg, deleted=True))
    return d


def _apply_segments_to_paragraph(
    paragraph, segments: List[Dict[str, Any]], author: str, when: str, rev: _RevisionCounter
) -> None:
    """Attach a stream of typed segments to a python-docx paragraph.
    Text runs go through python-docx's ordinary API; `ins` / `del`
    segments hand-craft the OOXML because python-docx doesn't expose a
    high-level API for revisions."""
    p_element = paragraph._p
    for seg in segments:
        kind = seg.get("kind")
        if kind == "text":
            p_element.append(_make_run(seg))
        elif kind == "linebreak":
            p_element.append(_make_linebreak_run())
        elif kind == "ins":
            p_element.append(_make_ins(seg["runs"], author, when, rev))
        elif kind == "del":
            p_element.append(_make_del(seg["runs"], author, when, rev))


# ─── Block renderer ─────────────────────────────────────────────────────────
_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _apply_alignment(paragraph, block: Dict[str, Any]) -> None:
    tunes = block.get("tunes") or {}
    align = ((tunes.get("alignment") or {}).get("alignment")) or "left"
    paragraph.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)


def _render_block(doc, block: Dict[str, Any], author: str, when: str, rev: _RevisionCounter) -> None:
    """Attach one EditorJS block to the given docx Document. Unknown
    block types fall back to a plain paragraph carrying the block's
    `text` field, so unsupported blocks never silently drop content."""
    if not isinstance(block, dict):
        return
    btype = (block.get("type") or "paragraph").lower()
    data = block.get("data") or {}
    text = data.get("text") if isinstance(data.get("text"), str) else ""

    if btype == "header":
        level = max(1, min(6, int(data.get("level") or 2)))
        p = doc.add_paragraph()
        _apply_alignment(p, block)
        # Word maps Heading N styles onto its outline levels; python-docx
        # exposes them under those names.
        p.style = doc.styles[f"Heading {level}"]
        _apply_segments_to_paragraph(p, _parse_inline(text), author, when, rev)
        return

    if btype in ("paragraph", "quote"):
        p = doc.add_paragraph()
        _apply_alignment(p, block)
        if btype == "quote":
            p.style = doc.styles["Intense Quote"]
        _apply_segments_to_paragraph(p, _parse_inline(text), author, when, rev)
        return

    if btype == "list":
        style_name = "List Number" if (data.get("style") == "ordered") else "List Bullet"
        items = data.get("items") or []
        for item in items:
            item_text = item if isinstance(item, str) else (item.get("content") or item.get("text") or "")
            p = doc.add_paragraph(style=style_name)
            _apply_alignment(p, block)
            _apply_segments_to_paragraph(p, _parse_inline(item_text), author, when, rev)
        return

    if btype == "checklist":
        items = data.get("items") or []
        for item in items:
            checked = bool(item.get("checked")) if isinstance(item, dict) else False
            marker = "☑ " if checked else "☐ "
            item_text = item.get("text") if isinstance(item, dict) else str(item or "")
            p = doc.add_paragraph()
            _apply_alignment(p, block)
            p.add_run(marker)
            _apply_segments_to_paragraph(p, _parse_inline(item_text or ""), author, when, rev)
        return

    if btype == "delimiter":
        # A centered dot row — Word doesn't have a first-class "delimiter"
        # so we replicate what most authoring tools do for section
        # breaks in narrative documents.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("• • •")
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        return

    if btype == "table":
        content = data.get("content") or []
        if not content or not isinstance(content, list):
            return
        rows = len(content)
        cols = max((len(r) for r in content if isinstance(r, list)), default=0)
        if cols == 0:
            return
        tbl = doc.add_table(rows=rows, cols=cols)
        tbl.style = "Table Grid"
        with_headings = bool(data.get("withHeadings"))
        for ri, row in enumerate(content):
            if not isinstance(row, list):
                continue
            for ci, cell in enumerate(row):
                if ci >= cols:
                    break
                cell_text = cell if isinstance(cell, str) else ""
                tcell = tbl.cell(ri, ci)
                # First paragraph created by add_table is empty; overwrite it.
                tcell.text = ""
                p = tcell.paragraphs[0]
                if with_headings and ri == 0:
                    p.style = doc.styles["Heading 4"]
                _apply_segments_to_paragraph(p, _parse_inline(cell_text), author, when, rev)
        return

    if btype == "image":
        # We don't embed remote images (would require fetching), but we
        # preserve the URL as a caption paragraph so nothing is silently
        # dropped. The caller can swap in an inline image later if we
        # decide to support that.
        src = ((data.get("file") or {}).get("url")) or data.get("url") or ""
        caption = data.get("caption") or ""
        if src:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"[image: {src}]")
            r.italic = True
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _apply_segments_to_paragraph(p, _parse_inline(caption), author, when, rev)
        return

    # Fallback for unknown block types — never silently drop content.
    if text:
        p = doc.add_paragraph()
        _apply_alignment(p, block)
        _apply_segments_to_paragraph(p, _parse_inline(text), author, when, rev)


# ─── Headers / footers ──────────────────────────────────────────────────────
def _normalise_hf(raw: Any) -> Dict[int, List[Dict[str, Any]]]:
    """Flatten the response's header / footer container (which can
    arrive as a dict, a list of {pageNo, blocks}, or Mongoose's Map
    with a `.entries()` method) into a plain `{page_index: blocks}`
    mapping keyed by 0-based page index."""
    if raw is None:
        return {}
    out: Dict[int, List[Dict[str, Any]]] = {}
    if isinstance(raw, list):
        for entry in raw:
            if not isinstance(entry, dict):
                continue
            idx = entry.get("pageNo") if "pageNo" in entry else entry.get("page_no")
            blocks = entry.get("blocks") or []
            try:
                out[int(idx)] = list(blocks) if isinstance(blocks, list) else []
            except (TypeError, ValueError):
                continue
        return out
    if isinstance(raw, dict):
        for k, v in raw.items():
            try:
                idx = int(k)
            except (TypeError, ValueError):
                continue
            if isinstance(v, dict) and isinstance(v.get("blocks"), list):
                out[idx] = list(v["blocks"])
            elif isinstance(v, list):
                out[idx] = list(v)
        return out
    return out


def _fill_hf_from_blocks(container, blocks: List[Dict[str, Any]], author: str, when: str, rev: _RevisionCounter) -> None:
    """Populate a header or footer container (both look the same to
    python-docx) with our block renderer. python-docx creates a single
    empty paragraph in the header/footer by default; we clear it first
    so nothing overlaps."""
    if not blocks:
        return
    for p in list(container.paragraphs):
        p._element.getparent().remove(p._element)
    # Route header/footer blocks through the same renderer as the body so
    # `<ins>` / `<del>` marks survive there too. If the contract has
    # different headers on different pages, that's handled by adding one
    # section per page with `is_linked_to_previous = False`.
    dummy_doc = container  # for the renderer we just need `.add_paragraph`
    for block in blocks:
        _render_hf_block(container, block, author, when, rev)


def _render_hf_block(container, block: Dict[str, Any], author: str, when: str, rev: _RevisionCounter) -> None:
    """Cut-down variant of `_render_block` that writes into a header /
    footer container (not the main body). Headers/footers can hold
    paragraphs, images, and tables — everything else falls back to
    paragraph text."""
    if not isinstance(block, dict):
        return
    btype = (block.get("type") or "paragraph").lower()
    data = block.get("data") or {}
    text = data.get("text") if isinstance(data.get("text"), str) else ""
    if btype == "header":
        level = max(1, min(6, int(data.get("level") or 2)))
        p = container.add_paragraph()
        p.style = container.part.document.styles[f"Heading {level}"]
        _apply_alignment(p, block)
        _apply_segments_to_paragraph(p, _parse_inline(text), author, when, rev)
        return
    # Everything else lands as a paragraph — headers/footers rarely need
    # tables/lists, and if we hit an edge case a plain paragraph is
    # graceful degradation.
    p = container.add_paragraph()
    _apply_alignment(p, block)
    _apply_segments_to_paragraph(p, _parse_inline(text), author, when, rev)


# ─── Public entry point ─────────────────────────────────────────────────────
def render_contract_docx(
    pages: List[Dict[str, Any]],
    meta: Dict[str, Any],
) -> bytes:
    """Assemble a Track-Changes-enabled DOCX for the given contract and
    return the file bytes.

    Arguments
    ---------
    pages : list of dicts shaped `{ "pageNo": int, "blocks": [...] }`
        The per-page EditorJS block dumps. Sorted by `pageNo` on entry
        so the caller doesn't have to.
    meta : dict
        Required keys:
          - "title" (str)  — used for the Word document title property
                             and rendered as an H1 at the top.
          - "author" (str) — populates `w:author` on every `w:ins` /
                             `w:del`; typically the sender's full name
                             (or their email if no name).
        Optional keys:
          - "message" (str)  — cover blurb rendered above the body.
          - "date" (str, ISO)— used for `w:date` on revisions; defaults
                               to now in UTC.
          - "headers" / "footers" — `{ pageIndex: {blocks} }` maps for
                               per-page header / footer content.
          - "enable_track_changes" (bool, default True) — whether to
                                open the file with Track Changes ON so
                                the counterparty's further edits get
                                recorded too.
    """
    title = str(meta.get("title") or "Contract")
    author = str(meta.get("author") or "Sender")
    message = meta.get("message") or ""
    when = meta.get("date") or datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    enable_tc = bool(meta.get("enable_track_changes", True))
    headers = _normalise_hf(meta.get("headers"))
    footers = _normalise_hf(meta.get("footers"))

    rev = _RevisionCounter()
    doc = Document()

    # Metadata — title + author show up in Word's document properties
    # dialog and on the recent-files hover cards.
    doc.core_properties.title = title
    doc.core_properties.author = author

    # Sort pages so page 1 lands first even when the caller passes an
    # unordered list.
    sorted_pages = sorted(pages or [], key=lambda p: int(p.get("pageNo") or 0))

    # Force Word to open the document with Track Changes turned ON so
    # any FURTHER edits the counterparty makes (beyond accepting/rejecting
    # what the sender sent) also get tracked. Without this Word opens
    # in normal editing mode and the counterparty's straight-typed
    # changes disappear into the file without a revision marker.
    if enable_tc:
        settings = doc.settings.element
        track = OxmlElement("w:trackChanges")
        settings.append(track)

    # Cover — title + optional intro paragraph. Kept small so the actual
    # contract sits at the top of page 1.
    title_p = doc.add_paragraph()
    title_p.style = doc.styles["Title"]
    title_p.add_run(title)
    if message:
        p = doc.add_paragraph()
        p.style = doc.styles["Intense Quote"]
        p.add_run(str(message))

    # Attach per-page header + footer via the section headers/footers.
    # Word associates one hdr/ftr per section, so when we detect
    # per-page differences we insert a section break at the top of the
    # affected page and unlink its header from the previous section.
    section = doc.sections[0]

    # Common case: all pages share page-0's header/footer (or the doc
    # has a single set that isn't keyed by page index at all). Attach
    # them to section 0 and let Word repeat them across every page.
    default_header_blocks = headers.get(0) or (list(headers.values())[0] if headers else [])
    default_footer_blocks = footers.get(0) or (list(footers.values())[0] if footers else [])
    if default_header_blocks:
        _fill_hf_from_blocks(section.header, default_header_blocks, author, when, rev)
    if default_footer_blocks:
        _fill_hf_from_blocks(section.footer, default_footer_blocks, author, when, rev)

    # Emit the body. Page N (1-indexed for the user, 0-indexed in our
    # data) gets an explicit page-break-before once we're past page 1.
    for idx, page in enumerate(sorted_pages):
        page_no = int(page.get("pageNo") or idx)
        blocks = page.get("blocks") or []
        if idx > 0:
            # A `w:br` with `w:type="page"` inside a run is Word's
            # explicit page break marker.
            p = doc.add_paragraph()
            r = p.add_run()
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            r._r.append(br)

        # Per-page header/footer overrides — only if a specific per-page
        # override exists AND it differs from the default. Adding a new
        # section for every page bloats the file, so we only pay the
        # cost when it's actually different content.
        override_header = headers.get(page_no)
        override_footer = footers.get(page_no)
        if idx > 0 and (
            (override_header and override_header is not default_header_blocks)
            or (override_footer and override_footer is not default_footer_blocks)
        ):
            new_section = doc.add_section()
            new_section.header.is_linked_to_previous = False
            new_section.footer.is_linked_to_previous = False
            if override_header:
                _fill_hf_from_blocks(new_section.header, override_header, author, when, rev)
            if override_footer:
                _fill_hf_from_blocks(new_section.footer, override_footer, author, when, rev)

        for block in blocks:
            _render_block(doc, block, author, when, rev)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Patch-in-place: preserve original DOCX fidelity ───────────────────────
#
# When the counterparty uploaded a `.docx` (counterparty-initiated flow),
# rendering a brand-new Word file from EditorJS blocks loses formatting
# the sender didn't touch — fonts, exact list numbering, embedded images,
# text-box positioning, custom styles Word bakes in for enterprise
# templates. Enterprise CLMs (Ironclad, DocuSign Negotiate) sidestep this
# by editing the ORIGINAL `.docx` in place: for each paragraph, they
# leave the OOXML untouched when the sender made no changes, and splice
# in `<w:ins>` / `<w:del>` only where changes exist.
#
# We do the same. Correlation between sender's EditorJS blocks and
# original `<w:p>` elements is done by **normalized-text similarity**:
# for each block we compute the "original" text (block text with
# `<ins>` content removed, `<mark>` content kept — i.e. what the block
# looked like BEFORE the sender's redlines), and find the `<w:p>` in
# the original DOCX whose normalized text matches best. Paragraphs the
# sender left completely untouched match on-the-nose and are skipped
# (perfect fidelity). Paragraphs the sender edited get their runs
# replaced with a new sequence carrying `<w:ins>` / `<w:del>` at the
# right offsets. New paragraphs (no match) are synthesized wrapped in
# `<w:ins>`. Original paragraphs no longer referenced by any block are
# wrapped in `<w:del>`.


def _norm_text(s: str) -> str:
    """Normalize a paragraph's text for similarity comparison. The goal
    is: two strings that a human would read as identical should hash
    identical, regardless of Word / EditorJS / HTML-entity artefacts.

    We collapse the differences we've actually seen in the wild:

      • curly quotes / apostrophes → straight
      • non-breaking spaces / zero-width joiners → normal space / dropped
      • HTML entities (`&amp;`, `&nbsp;`) → their real characters
      • em-dash / en-dash → simple hyphen
      • all whitespace runs → single space
      • trim + lowercase
    """
    if not s:
        return ""
    # Unescape HTML entities Word might have carried in from a paste.
    try:
        import html as _html
        s = _html.unescape(s)
    except Exception:
        pass
    # Straighten quotes and apostrophes.
    s = s.replace("‘", "'").replace("’", "'").replace("‚", "'").replace("‛", "'")
    s = s.replace("“", '"').replace("”", '"').replace("„", '"').replace("‟", '"')
    # Dashes → hyphen.
    s = s.replace("–", "-").replace("—", "-").replace("−", "-")
    # Bullet glyphs → drop (list items reach this function without <li>
    # markers, but Word occasionally paints a bullet as literal text).
    s = s.replace("•", "").replace("·", "").replace("◦", "")
    # Zero-width joiners / non-joiners / BOM → drop.
    s = re.sub(r"[​‌‍﻿]", "", s)
    # Non-breaking space / narrow NBSP → regular space.
    s = s.replace(" ", " ").replace(" ", " ")
    # Collapse any run of whitespace to one space.
    s = re.sub(r"\s+", " ", s).strip().lower()
    return s


def _visible_text_from_html(html: str) -> str:
    """Strip HTML but preserve visible text — used for the block's
    'after-edit' representation when we need to compare against the
    original paragraph's plain text. Keeps insertion content, drops
    deletion content (the sender said 'delete this')."""
    if not html:
        return ""
    return _keep_ins_strip_mark(html)


def _wp_text(wp) -> str:
    """Extract the visible text out of a `<w:p>` element. Concatenates
    `<w:t>` and `<w:delText>` content across all runs in the paragraph,
    including runs nested inside existing `<w:ins>` / `<w:del>` (though
    a fresh original DOCX shouldn't have those)."""
    parts: List[str] = []
    for t in wp.iter():
        if t.tag == qn("w:t") or t.tag == qn("w:delText"):
            parts.append(t.text or "")
        elif t.tag == qn("w:br"):
            parts.append("\n")
        elif t.tag == qn("w:tab"):
            parts.append("\t")
    return "".join(parts)


def _strip_ins_keep_mark(html: str) -> str:
    """Return the block's text as it looked BEFORE the sender's edits.
    Sender inserts (`<ins class="clm-suggestion-insert">`) → removed.
    Sender deletions (`<mark class="clm-redline">`) → kept (they existed
    in the original). Everything else falls through as plain text."""
    if not html:
        return ""
    soup = BeautifulSoup(html, "html.parser")
    for ins in soup.find_all("ins", class_=_TRACKED_INS_CLASS):
        ins.decompose()
    return soup.get_text()


def _keep_ins_strip_mark(html: str) -> str:
    """Return the block's text as it looks AFTER applying the sender's
    edits. Inserts kept, deletions removed. Used to detect no-op
    paragraphs (original == new = no edits at all)."""
    if not html:
        return ""
    soup = BeautifulSoup(html, "html.parser")
    for m in soup.find_all(["mark", "del"], class_=_TRACKED_DEL_CLASS):
        m.decompose()
    return soup.get_text()


def _block_visible_html(block: Dict[str, Any]) -> str:
    """Pull the block's inline HTML for text/header/quote/list-item.
    Non-inline block types (image, table, delimiter) return empty."""
    if not isinstance(block, dict):
        return ""
    data = block.get("data") or {}
    text = data.get("text")
    if isinstance(text, str):
        return text
    return ""


def _build_pid_index(
    body,
) -> Dict[str, List[Any]]:
    """Walk the document body and build `{ normalized_text: [<w:p>, ...] }`.
    Duplicate paragraphs (empty, page-repeating footers) map to a list;
    the matcher pops from the front so multiple identical rows still
    correlate 1:1 with successive blocks."""
    index: Dict[str, List[Any]] = {}
    for child in body:
        if child.tag != qn("w:p"):
            continue
        key = _norm_text(_wp_text(child))
        index.setdefault(key, []).append(child)
    return index


def _clear_wp_runs(wp) -> None:
    """Remove every child of `<w:p>` except `<w:pPr>`. Leaves paragraph
    properties (style, alignment, numbering) intact so replacing the
    run contents keeps the paragraph looking the same."""
    to_remove = [c for c in wp if c.tag != qn("w:pPr")]
    for c in to_remove:
        wp.remove(c)


def _append_segments_to_wp(
    wp, segments: List[Dict[str, Any]], author: str, when: str, rev: _RevisionCounter
) -> None:
    """Attach a stream of typed segments (from _parse_inline) directly
    to a raw `<w:p>` element. Mirrors _apply_segments_to_paragraph but
    without the python-docx paragraph wrapper — we already own the
    element from the original DOCX and need to append into it directly."""
    for seg in segments:
        kind = seg.get("kind")
        if kind == "text":
            wp.append(_make_run(seg))
        elif kind == "linebreak":
            wp.append(_make_linebreak_run())
        elif kind == "ins":
            wp.append(_make_ins(seg["runs"], author, when, rev))
        elif kind == "del":
            wp.append(_make_del(seg["runs"], author, when, rev))


def _wrap_wp_in_del(wp, author: str, when: str, rev: _RevisionCounter) -> None:
    """Mark a paragraph the sender removed. We can't wrap the whole
    `<w:p>` in `<w:del>` (invalid OOXML), so we walk its runs, convert
    each `<w:t>` inside to a `<w:delText>`, and wrap the runs in a
    single `<w:del>` element. The paragraph properties are preserved."""
    # Collect every run's text content and its rPr.
    runs_to_delete: List[Any] = []
    for child in list(wp):
        if child.tag == qn("w:r"):
            runs_to_delete.append(child)
            wp.remove(child)
    if not runs_to_delete:
        return
    # Convert `<w:t>` in each run to `<w:delText>`.
    for r in runs_to_delete:
        for t in list(r):
            if t.tag == qn("w:t"):
                new_t = OxmlElement("w:delText")
                new_t.text = t.text or ""
                if t.get(qn("xml:space")):
                    new_t.set(qn("xml:space"), t.get(qn("xml:space")))
                else:
                    new_t.set(qn("xml:space"), "preserve")
                r.remove(t)
                r.append(new_t)
    d = OxmlElement("w:del")
    d.set(qn("w:id"), rev.next())
    d.set(qn("w:author"), author)
    d.set(qn("w:date"), when)
    for r in runs_to_delete:
        d.append(r)
    wp.append(d)


def _make_new_wp_wrapped_in_ins(
    block: Dict[str, Any], author: str, when: str, rev: _RevisionCounter
) -> Any:
    """Synthesize a brand-new `<w:p>` for a paragraph the sender added
    entirely. Runs go inside `<w:ins>` so the whole paragraph shows as
    an insertion in the Review ribbon."""
    wp = OxmlElement("w:p")
    text = _block_visible_html(block)
    segments = _parse_inline(text)
    # Everything in a brand-new paragraph is a tracked insertion, so
    # wrap the entire segment sequence in a single `<w:ins>`.
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), rev.next())
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), when)
    for seg in segments:
        kind = seg.get("kind")
        if kind == "text":
            ins.append(_make_run(seg))
        elif kind == "linebreak":
            ins.append(_make_linebreak_run())
        elif kind == "ins":
            for r_seg in seg.get("runs") or []:
                if r_seg.get("kind") == "linebreak":
                    ins.append(_make_linebreak_run())
                else:
                    ins.append(_make_run(r_seg))
        elif kind == "del":
            # A del inside a brand-new paragraph is odd but not
            # impossible (sender re-edited their own insert). Emit
            # a nested `<w:del>` inside the outer `<w:ins>`.
            d = OxmlElement("w:del")
            d.set(qn("w:id"), rev.next())
            d.set(qn("w:author"), author)
            d.set(qn("w:date"), when)
            for r_seg in seg.get("runs") or []:
                d.append(_make_run(r_seg, deleted=True))
            ins.append(d)
    wp.append(ins)
    return wp


def _find_matching_wp(
    original_key: str,
    index: Dict[str, List[Any]],
) -> Optional[Any]:
    """Pop the first `<w:p>` matching this normalized text. Empty keys
    don't match — an empty original block usually means the sender
    inserted a paragraph from nothing, which the caller handles with
    _make_new_wp_wrapped_in_ins."""
    if not original_key:
        return None
    bucket = index.get(original_key)
    if not bucket:
        return None
    return bucket.pop(0)


def _find_matching_wp_by_prefix(
    candidate: str,
    index: Dict[str, List[Any]],
    min_words: int = 6,
) -> Optional[Any]:
    """Fallback correlation. When exact match fails, look for a
    paragraph that shares its first `min_words` words with the block's
    current text. Catches the direct-typing case where the sender
    changed the middle of a paragraph — the prefix still lines up so we
    can locate the source paragraph and diff it.

    Iterates over all remaining un-popped buckets. First hit wins
    (stable order in Python 3.7+, so this walks the DOCX in reading
    order — matching what a human would expect)."""
    if not candidate:
        return None
    words = candidate.split(" ")
    if len(words) < min_words:
        return None
    prefix = " ".join(words[:min_words])
    best_key: Optional[str] = None
    for key, bucket in index.items():
        if not bucket or not key:
            continue
        key_words = key.split(" ")
        if len(key_words) < min_words:
            continue
        key_prefix = " ".join(key_words[:min_words])
        if key_prefix == prefix:
            best_key = key
            break
    if best_key is None:
        return None
    return index[best_key].pop(0)


def _find_matching_wp_by_similarity(
    candidate: str,
    index: Dict[str, List[Any]],
    min_ratio: float = 0.55,
) -> Optional[Any]:
    """Best-effort fuzzy correlation. When both exact and prefix match
    fail, walk the remaining un-popped keys and pick the one with the
    highest SequenceMatcher ratio against `candidate`. Threshold of
    0.55 catches "changed a phrase in the middle of a paragraph"
    without pairing unrelated short paragraphs by coincidence.

    Guardrails:
      • Very short candidates (fewer than 4 words) skip the fuzzy
        search — too likely to false-positive on headings.
      • The single-highest match is chosen; ties break on document
        order via dict iteration order (Python 3.7+ insertion order).
    """
    from difflib import SequenceMatcher

    if not candidate:
        return None
    if len(candidate.split(" ")) < 4:
        return None
    best_key: Optional[str] = None
    best_ratio: float = 0.0
    for key, bucket in index.items():
        if not bucket or not key:
            continue
        r = SequenceMatcher(a=candidate, b=key, autojunk=False).ratio()
        if r > best_ratio:
            best_ratio = r
            best_key = key
    if best_key is None or best_ratio < min_ratio:
        return None
    return index[best_key].pop(0)


def _tokenize_for_diff(text: str) -> List[str]:
    """Split text into diffable tokens. Words and punctuation are kept
    as separate tokens so a single-word edit produces one ins + one del
    instead of a whole-sentence replacement. Whitespace becomes its own
    token so we can preserve spacing on reassembly."""
    if not text:
        return []
    # Match word runs, whitespace runs, or single non-word non-space characters.
    return re.findall(r"\w+|\s+|[^\w\s]", text, flags=re.UNICODE)


def _append_diff_segments_to_wp(
    wp, original_plain: str, new_plain: str,
    author: str, when: str, rev: _RevisionCounter,
) -> None:
    """Compute a word-level diff between the original paragraph text
    and the sender's (post-edit) text, then emit runs into `wp` where
    unchanged tokens are plain `<w:r>`, added tokens are wrapped in
    `<w:ins>`, and removed tokens are wrapped in `<w:del>`. This is
    the "no suggestion mode" fallback — the sender edited the text
    directly, and we synthesise Track Changes from the diff."""
    from difflib import SequenceMatcher

    a = _tokenize_for_diff(original_plain)
    b = _tokenize_for_diff(new_plain)
    sm = SequenceMatcher(a=a, b=b, autojunk=False)

    def _emit_plain(tokens: List[str]) -> None:
        if not tokens:
            return
        text = "".join(tokens)
        wp.append(_make_run({"kind": "text", "text": text}))

    def _emit_ins(tokens: List[str]) -> None:
        if not tokens:
            return
        text = "".join(tokens)
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), rev.next())
        ins.set(qn("w:author"), author)
        ins.set(qn("w:date"), when)
        ins.append(_make_run({"kind": "text", "text": text}))
        wp.append(ins)

    def _emit_del(tokens: List[str]) -> None:
        if not tokens:
            return
        text = "".join(tokens)
        d = OxmlElement("w:del")
        d.set(qn("w:id"), rev.next())
        d.set(qn("w:author"), author)
        d.set(qn("w:date"), when)
        d.append(_make_run({"kind": "text", "text": text}, deleted=True))
        wp.append(d)

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            _emit_plain(a[i1:i2])
        elif tag == "delete":
            _emit_del(a[i1:i2])
        elif tag == "insert":
            _emit_ins(b[j1:j2])
        elif tag == "replace":
            # Emit deletion FIRST then insertion — matches how Word
            # itself serialises simultaneous edits on the same range.
            _emit_del(a[i1:i2])
            _emit_ins(b[j1:j2])


def _has_any_edits(block: Dict[str, Any]) -> bool:
    """Cheap check: does the block's HTML contain any tracked-change
    marks (`<ins class="clm-suggestion-insert">` or `<mark class="clm-redline">`)?
    If not, the paragraph is untouched and we can skip patching entirely
    — the original `<w:p>` stays byte-identical."""
    html = _block_visible_html(block)
    if not html:
        return False
    if _TRACKED_INS_CLASS in html or _TRACKED_DEL_CLASS in html:
        return True
    return False


def patch_contract_docx(
    original_bytes: bytes,
    pages: List[Dict[str, Any]],
    meta: Dict[str, Any],
) -> bytes:
    """Enterprise-standard round-trip. Open the counterparty's original
    `.docx`, splice in the sender's `<w:ins>` / `<w:del>` at matching
    paragraphs, and return the modified file. Paragraphs the sender
    didn't touch stay pixel-identical; the whole rest of the DOCX
    (headers, footers, images, styles, section properties) is
    preserved verbatim.
    """
    author = str(meta.get("author") or "Sender")
    when = meta.get("date") or datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    enable_tc = bool(meta.get("enable_track_changes", True))

    doc = Document(io.BytesIO(original_bytes))
    body = doc.element.body

    # Force Track Changes ON so the counterparty's further edits are
    # also recorded when they open the file.
    if enable_tc:
        settings = doc.settings.element
        # Don't add w:trackChanges twice if it's already there.
        if settings.find(qn("w:trackChanges")) is None:
            settings.append(OxmlElement("w:trackChanges"))

    rev = _RevisionCounter()

    # Flatten pages into one ordered list — the original DOCX doesn't
    # know about EditorJS page boundaries.
    ordered_blocks: List[Dict[str, Any]] = []
    for page in sorted(pages or [], key=lambda p: int(p.get("pageNo") or 0)):
        for b in page.get("blocks") or []:
            if isinstance(b, dict):
                ordered_blocks.append(b)

    # Index the body's paragraphs for matching. We use a fresh index
    # per call because pop() mutates the buckets.
    pid_index = _build_pid_index(body)

    # Walk blocks in order. For each, decide: skip (no edits) / patch
    # (edited existing) / synthesize (new). Record which <w:p> each
    # block landed on so remaining unmatched <w:p>s can be marked
    # deleted at the end.
    matched_wps: set = set()
    # For new-paragraph insertions we need an anchor to splice after —
    # the last matched original paragraph.
    last_anchor: Optional[Any] = None

    # Flatten list-item + checklist-item blocks into pseudo-blocks so
    # each item correlates 1:1 with its own <w:p> in the original DOCX
    # (that's how Word stores lists — one paragraph per item, with a
    # `<w:numPr>` binding them to a numbering definition). Without this
    # every list item became a "new insert" and every original list
    # paragraph became "deleted", turning the whole list into a
    # revision block.
    flat_blocks: List[Dict[str, Any]] = []
    for block in ordered_blocks:
        btype = (block.get("type") or "paragraph").lower()
        data = block.get("data") or {}
        if btype in ("list", "checklist"):
            items = data.get("items") or []
            for item in items:
                # Items are either strings ("hello world") or dicts
                # with `content`/`text` fields (EditorJS's newer list
                # tool nests). Both reach here as a paragraph pseudo-
                # block so downstream matching works uniformly.
                item_text = ""
                if isinstance(item, str):
                    item_text = item
                elif isinstance(item, dict):
                    item_text = (
                        item.get("content")
                        or item.get("text")
                        or ""
                    )
                flat_blocks.append({
                    "type": "paragraph",
                    "data": {"text": item_text},
                    "_from_list": True,
                })
            continue
        flat_blocks.append(block)

    for block in flat_blocks:
        btype = (block.get("type") or "paragraph").lower()
        # Table / image / delimiter blocks don't map to a single <w:p>.
        # Skip the correlation attempt so we don't accidentally match
        # a stray paragraph and clobber unrelated content. If the
        # sender edited a table, the current patch algorithm won't
        # carry it over — that's a known limitation flagged in the
        # spec and would require a table-aware diff to fix.
        if btype not in ("paragraph", "header", "quote"):
            continue

        html = _block_visible_html(block)
        # "before" and "after" text for correlation:
        # - original_text = text as it looked BEFORE the sender's edits
        #   (strip <ins>, keep <mark>). This is what matches the source
        #   <w:p> in the original DOCX.
        # - new_text = text as it looks AFTER the sender's edits
        #   (keep <ins>, strip <mark>). This is what should end up in
        #   the DOCX after patching.
        original_text = _norm_text(_strip_ins_keep_mark(html))
        new_text = _norm_text(_keep_ins_strip_mark(html))

        # Direct-typing case: the sender edited the block WITHOUT using
        # suggestion mode (no <ins>/<mark>) — so `original_text` and
        # `new_text` are BOTH the current block text. In that case
        # we can't match by "before" text (the block has no memory of
        # its pre-edit state). Correlate by AFTER text and, if that
        # matches an original paragraph exactly, treat it as untouched.
        # If it DOESN'T match, we fall back to correlating by prefix
        # (first ~12 words) so a modest edit still lands on the right
        # paragraph.
        has_marks = _has_any_edits(block)
        match = _find_matching_wp(original_text, pid_index)

        if match is None and not has_marks:
            # No explicit marks + no exact match on the current text.
            # Try prefix first (cheap, high-precision), then fall back
            # to fuzzy similarity so an edit that touched the middle of
            # a paragraph still correlates to its source <w:p>. Without
            # this, direct-typed edits get treated as "brand-new
            # paragraph" + the original as "deleted", which is exactly
            # the "changes don't match what I see in the editor"
            # symptom in production.
            match = (
                _find_matching_wp_by_prefix(new_text, pid_index)
                or _find_matching_wp_by_similarity(new_text, pid_index)
            )
        if match is None and has_marks:
            # Explicit marks but no exact match on before-text — likely
            # a spelling difference between the DOCX and how EditorJS
            # rendered it. Same prefix→fuzzy waterfall.
            match = (
                _find_matching_wp_by_prefix(original_text, pid_index)
                or _find_matching_wp_by_similarity(original_text, pid_index)
            )

        if match is not None:
            matched_wps.add(id(match))
            last_anchor = match

            wp_text_now = _norm_text(_wp_text(match))
            no_edits_at_all = (
                not has_marks and _norm_text(new_text) == wp_text_now
            )
            if no_edits_at_all:
                # Untouched paragraph — leave <w:p> exactly as-is.
                continue

            if has_marks:
                # Sender used suggestion mode. Replace runs with the
                # explicit <ins>/<del> sequence.
                _clear_wp_runs(match)
                segments = _parse_inline(html)
                _append_segments_to_wp(match, segments, author, when, rev)
            else:
                # Direct typing without suggestion mode. Compute a
                # word-level diff between the original paragraph and
                # the new block text so we still get tracked changes
                # (this is what enterprise CLMs do with Word's
                # "Compare Documents" feature). Extract the original
                # text BEFORE clearing runs — otherwise the diff sees
                # an empty "before" side.
                orig_plain = _wp_text(match)
                new_plain = _visible_text_from_html(html)
                _clear_wp_runs(match)
                _append_diff_segments_to_wp(
                    match, orig_plain, new_plain, author, when, rev,
                )
        else:
            # No match anywhere — treat as a brand-new paragraph the
            # sender inserted from scratch. Synthesize <w:p> wrapped in
            # <w:ins> and splice after the last anchor.
            new_wp = _make_new_wp_wrapped_in_ins(block, author, when, rev)
            if last_anchor is not None:
                last_anchor.addnext(new_wp)
                last_anchor = new_wp
            else:
                # No anchor yet — prepend before the first existing <w:p>.
                first_p = None
                for child in body:
                    if child.tag == qn("w:p"):
                        first_p = child
                        break
                if first_p is not None:
                    first_p.addprevious(new_wp)
                else:
                    body.append(new_wp)
                last_anchor = new_wp

    # Any <w:p> still sitting in pid_index (unpopped) was NOT referenced
    # by any block — the sender removed it. Wrap in <w:del>. We skip
    # completely-empty paragraphs (their normalized text is "") because
    # those exist as blank lines between paragraphs and the sender's
    # block list usually doesn't include them.
    for key, bucket in pid_index.items():
        if not key:
            continue
        for wp in bucket:
            if id(wp) in matched_wps:
                continue
            _wrap_wp_in_del(wp, author, when, rev)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = ["render_contract_docx", "patch_contract_docx"]
